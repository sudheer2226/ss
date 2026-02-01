
# from flask import Flask, request, render_template, redirect, url_for, session, jsonify, send_file
# from werkzeug.security import generate_password_hash, check_password_hash
# from werkzeug.utils import secure_filename
# import sqlite3
# import os
# import io
# import PyPDF2
# import docx
# from google import genai

# # ---------------- CONFIG ----------------
# API_KEY = "AIzaSyBbN6dYdLdUmstSTnTZcBGQkzOI8iIAhwg"
# MODEL_NAME = "models/gemini-2.5-flash"
# CHUNK_SIZE = 2000

# LANGUAGE_MAP = {
#     "en": "English",
#     "hi": "Hindi",
#     "te": "Telugu",
#     "ta": "Tamil",
#     "kn": "Kannada",
#     "ml": "Malayalam"
# }

# UPLOAD_FOLDER = "uploads"
# os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# client = genai.Client(api_key=API_KEY)
# app = Flask(__name__)
# app.secret_key = "supersecretkey"
# app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# # ---------------- DATABASE ----------------
# def init_db():
#     conn = sqlite3.connect("users.db")
#     c = conn.cursor()
#     c.execute("""CREATE TABLE IF NOT EXISTS users (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         username TEXT UNIQUE,
#         password TEXT
#     )""")
#     conn.commit()
#     conn.close()

# init_db()

# def add_user(username, password):
#     conn = sqlite3.connect("users.db")
#     c = conn.cursor()
#     hashed = generate_password_hash(password)
#     try:
#         c.execute("INSERT INTO users (username, password) VALUES (?, ?)", (username, hashed))
#         conn.commit()
#         return True
#     except sqlite3.IntegrityError:
#         return False
#     finally:
#         conn.close()

# def verify_user(username, password):
#     conn = sqlite3.connect("users.db")
#     c = conn.cursor()
#     c.execute("SELECT password FROM users WHERE username=?", (username,))
#     row = c.fetchone()
#     conn.close()
#     if row and check_password_hash(row[0], password):
#         return True
#     return False

# # ---------------- FILE READERS ----------------
# def read_pdf(file):
#     text = ""
#     reader = PyPDF2.PdfReader(file)
#     for page in reader.pages:
#         extracted = page.extract_text()
#         if extracted:
#             text += extracted + "\n"
#     return text

# def read_docx(file):
#     document = docx.Document(file)
#     return "\n".join(p.text for p in document.paragraphs)

# def read_txt(file):
#     return file.read().decode("utf-8")

# def extract_text(uploaded_file):
#     name = uploaded_file.filename.lower()
#     if name.endswith(".pdf"):
#         return read_pdf(uploaded_file)
#     elif name.endswith(".docx"):
#         return read_docx(uploaded_file)
#     elif name.endswith(".txt"):
#         return read_txt(uploaded_file)
#     else:
#         raise ValueError("Unsupported file format")

# # ---------------- PROMPT BUILDERS ----------------
# def build_analysis_prompt(contract_text, language):
#     return f"""
# You are an AI legal contract analysis assistant.
# Analyze the contract below STRICTLY in {language}.
# Explain everything in SIMPLE, EASY language so a non-lawyer can understand.
# Follow this structure:
# 1. Explanation of the contract
# 2. Risky clauses
# 3. Obligations of each party
# 4. Payment, termination, and liability clauses
# Do NOT invent facts.
# CONTRACT TEXT:
# {contract_text}
# """

# def build_chat_prompt(contract_text, user_question, language):
#     return f"""
# You are an AI legal assistant.
# Contract:
# {contract_text}
# User Question: {user_question}
# Answer in {language} in very simple, clear language. Explain like a friendly guide.
# """

# # ---------------- AI ANALYSIS ----------------
# def analyze_contract(contract_text, language):
#     chunks = [contract_text[i:i+CHUNK_SIZE] for i in range(0, len(contract_text), CHUNK_SIZE)]
#     results = []
#     for chunk in chunks:
#         prompt = build_analysis_prompt(chunk, language)
#         response = client.models.generate_content(
#             model=MODEL_NAME,
#             contents=prompt
#         )
#         results.append(response.text)
#     return "\n\n".join(results)

# def chat_with_ai(contract_text, question, language):
#     prompt = build_chat_prompt(contract_text, question, language)
#     response = client.models.generate_content(
#         model=MODEL_NAME,
#         contents=prompt
#     )
#     return response.text

# # ---------------- ROUTES ----------------
# @app.route("/")
# def home():
#     if "username" in session:
#         return redirect(url_for("index"))
#     return redirect(url_for("login"))

# @app.route("/register", methods=["GET", "POST"])
# def register():
#     message = ""
#     if request.method == "POST":
#         username = request.form["username"]
#         password = request.form["password"]
#         if add_user(username, password):
#             return redirect(url_for("login"))
#         else:
#             message = "Username already exists."
#     return render_template("register.html", message=message)

# @app.route("/login", methods=["GET", "POST"])
# def login():
#     message = ""
#     if request.method == "POST":
#         username = request.form["username"]
#         password = request.form["password"]
#         if verify_user(username, password):
#             session["username"] = username
#             return redirect(url_for("index"))
#         else:
#             message = "Invalid username or password."
#     return render_template("login.html", message=message)

# @app.route("/logout")
# def logout():
#     session.pop("username", None)
#     return redirect(url_for("login"))

# @app.route("/index", methods=["GET", "POST"])
# def index():
#     if "username" not in session:
#         return redirect(url_for("login"))

#     analysis = ""
#     contract_text = ""
#     selected_language = "en"
#     if request.method == "POST":
#         uploaded_file = request.files.get("file")
#         selected_language = request.form.get("language", "en")
#         if not uploaded_file:
#             analysis = "Error: No file uploaded."
#         else:
#             try:
#                 contract_text = extract_text(uploaded_file)
#                 analysis = analyze_contract(contract_text, LANGUAGE_MAP[selected_language])
#             except Exception as e:
#                 analysis = f"Error: {str(e)}"

#     return render_template(
#         "index.html",
#         analysis=analysis,
#         languages=LANGUAGE_MAP,
#         selected_language=selected_language,
#         contract_text=contract_text
#     )

# @app.route("/download", methods=["POST"])
# def download():
#     if "username" not in session:
#         return redirect(url_for("login"))
#     content = request.form.get("analysis_content", "")
#     if not content:
#         return redirect(url_for("index"))
#     return send_file(
#         io.BytesIO(content.encode("utf-8")),
#         download_name="contract_analysis.txt",
#         as_attachment=True
#     )

# @app.route("/ask", methods=["POST"])
# def ask():
#     data = request.get_json()
#     contract_text = data.get("contract_text", "")
#     question = data.get("question", "")
#     language_code = data.get("language", "en")
#     language = LANGUAGE_MAP.get(language_code, "English")

#     if not contract_text or not question:
#         return jsonify({"answer": "Please upload a contract and ask a valid question."})
#     try:
#         answer = chat_with_ai(contract_text, question, language)
#         return jsonify({"answer": answer})
#     except Exception as e:
#         return jsonify({"answer": f"Error: {str(e)}"})

# if __name__ == "__main__":
#     app.run(debug=True)









from flask import Flask, request, render_template, redirect, url_for, session, jsonify, send_file
from werkzeug.security import generate_password_hash, check_password_hash
import sqlite3
import os
import io
import PyPDF2
import docx
import time
from google import genai

# ---------------- CONFIG ----------------
API_KEY = "AIzaSyCxMWULIZCmKLCPTj8zu4W8Ergk6zH8k30"   # 🔒 move to env variable later
MODEL_NAME = "models/gemini-2.5-flash"  # ✅ safer for free tier
MAX_TEXT_LENGTH = 12000         # ✅ single-call safe limit

LANGUAGE_MAP = {
    "en": "English",
    "hi": "Hindi",
    "te": "Telugu",
    "ta": "Tamil",
    "kn": "Kannada",
    "ml": "Malayalam"
}

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

client = genai.Client(api_key=API_KEY)
app = Flask(__name__)
app.secret_key = "supersecretkey"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# ---------------- DATABASE ----------------
def init_db():
    conn = sqlite3.connect("users.db")
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE,
            password TEXT
        )
    """)
    conn.commit()
    conn.close()

init_db()

def add_user(username, password):
    conn = sqlite3.connect("users.db")
    c = conn.cursor()
    hashed = generate_password_hash(password)
    try:
        c.execute("INSERT INTO users VALUES (NULL, ?, ?)", (username, hashed))
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        return False
    finally:
        conn.close()

def verify_user(username, password):
    conn = sqlite3.connect("users.db")
    c = conn.cursor()
    c.execute("SELECT password FROM users WHERE username=?", (username,))
    row = c.fetchone()
    conn.close()
    return row and check_password_hash(row[0], password)

# ---------------- FILE READERS ----------------
def read_pdf(file):
    reader = PyPDF2.PdfReader(file)
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def read_docx(file):
    document = docx.Document(file)
    return "\n".join(p.text for p in document.paragraphs)

def read_txt(file):
    return file.read().decode("utf-8")

def extract_text(uploaded_file):
    name = uploaded_file.filename.lower()
    if name.endswith(".pdf"):
        return read_pdf(uploaded_file)
    elif name.endswith(".docx"):
        return read_docx(uploaded_file)
    elif name.endswith(".txt"):
        return read_txt(uploaded_file)
    else:
        raise ValueError("Unsupported file format")

# ---------------- PROMPTS ----------------
def build_analysis_prompt(contract_text, language):
    return f"""
You are an AI legal contract analysis assistant.
Explain in SIMPLE {language}.
Structure:
1. Contract overview
2. Risky clauses
3. Responsibilities of each party
4. Payment, termination & liability
Do NOT invent information.

CONTRACT:
{contract_text}
"""

def build_chat_prompt(contract_text, question, language):
    return f"""
You are a helpful legal assistant.
Answer in SIMPLE {language}.

CONTRACT:
{contract_text}

QUESTION:
{question}
"""

# ---------------- SAFE GEMINI CALL ----------------
def safe_generate(prompt):
    for _ in range(3):
        try:
            return client.models.generate_content(
                model=MODEL_NAME,
                contents=prompt
            ).text
        except Exception as e:
            if "429" in str(e):
                time.sleep(20)
            else:
                raise e
    return "⚠️ AI is busy. Please try again after some time."

# ---------------- AI FUNCTIONS ----------------
def analyze_contract(contract_text, language):
    trimmed_text = contract_text[:MAX_TEXT_LENGTH]
    prompt = build_analysis_prompt(trimmed_text, language)
    return safe_generate(prompt)

def chat_with_ai(contract_text, question, language):
    trimmed_text = contract_text[:MAX_TEXT_LENGTH]
    prompt = build_chat_prompt(trimmed_text, question, language)
    return safe_generate(prompt)

# ---------------- ROUTES ----------------
@app.route("/")
def home():
    return redirect(url_for("index")) if "username" in session else redirect(url_for("login"))

@app.route("/register", methods=["GET", "POST"])
def register():
    msg = ""
    if request.method == "POST":
        if add_user(request.form["username"], request.form["password"]):
            return redirect(url_for("login"))
        msg = "Username already exists"
    return render_template("register.html", message=msg)

@app.route("/login", methods=["GET", "POST"])
def login():
    msg = ""
    if request.method == "POST":
        if verify_user(request.form["username"], request.form["password"]):
            session["username"] = request.form["username"]
            return redirect(url_for("index"))
        msg = "Invalid credentials"
    return render_template("login.html", message=msg)
@app.route("/about")
def about():
    return render_template("about.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

@app.route("/index", methods=["GET", "POST"])
def index():
    if "username" not in session:
        return redirect(url_for("login"))

    analysis = ""
    contract_text = ""
    selected_language = "en"

    if request.method == "POST":
        file = request.files.get("file")
        selected_language = request.form.get("language", "en")
        if not file:
            analysis = "No file uploaded"
        else:
            try:
                contract_text = extract_text(file)
                analysis = analyze_contract(contract_text, LANGUAGE_MAP[selected_language])
            except Exception as e:
                analysis = str(e)

    return render_template(
        "index.html",
        analysis=analysis,
        languages=LANGUAGE_MAP,
        selected_language=selected_language,
        contract_text=contract_text
    )

@app.route("/download", methods=["POST"])
def download():
    content = request.form.get("analysis_content", "")
    return send_file(
        io.BytesIO(content.encode()),
        download_name="contract_analysis.txt",
        as_attachment=True
    )

@app.route("/ask", methods=["POST"])
def ask():
    data = request.get_json()
    answer = chat_with_ai(
        data.get("contract_text", ""),
        data.get("question", ""),
        LANGUAGE_MAP.get(data.get("language", "en"), "English")
    )
    return jsonify({"answer": answer})

# ---------------- RUN ----------------
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)

